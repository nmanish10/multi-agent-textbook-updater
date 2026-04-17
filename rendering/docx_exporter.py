from __future__ import annotations

from pathlib import Path
from typing import List, Tuple

from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Inches, Pt, RGBColor

from core.models import Book, WrittenUpdate
from rendering.export_utils import build_export_manifest, write_export_manifest
from rendering.markdown_renderer import assign_subsection_ids


def _source_lines(update: WrittenUpdate) -> List[str]:
    if update.sources:
        source = update.sources[0]
        lines = []
        if source.title and source.date:
            lines.append(f"{source.title} ({source.date})")
        elif source.title:
            lines.append(source.title)
        if source.url:
            lines.append(source.url)
        return lines
    return [update.source] if update.source else []


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    styles["Normal"].font.name = "Georgia"
    styles["Normal"].font.size = Pt(11)

    for style_name, size, color in [
        ("Title", 22, RGBColor(0x17, 0x32, 0x4D)),
        ("Heading 1", 18, RGBColor(0x17, 0x32, 0x4D)),
        ("Heading 2", 15, RGBColor(0x1C, 0x5D, 0x7F)),
        ("Heading 3", 12, RGBColor(0x22, 0x48, 0x5F)),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = color


def _add_title_page(document: Document, book_title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(book_title)
    run.bold = True
    run.font.size = Pt(26)
    run.font.name = "Aptos Display"
    run.font.color.rgb = RGBColor(0x17, 0x32, 0x4D)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Updated Textbook Export")
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x6B, 0x7A)

    document.add_paragraph()
    document.add_section(WD_SECTION.NEW_PAGE)


def _add_export_summary(document: Document, book: Book, updates: List[WrittenUpdate]) -> None:
    document.add_heading("Export Summary", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Field"
    header_cells[1].text = "Value"
    _set_cell_shading(header_cells[0], "D9E7F0")
    _set_cell_shading(header_cells[1], "D9E7F0")

    fields = [
        ("Chapters", str(len(book.chapters))),
        ("Sections", str(sum(len(chapter.sections) for chapter in book.chapters))),
        ("Accepted updates", str(len(updates))),
        ("Canonical source", "Markdown-first publishing pipeline"),
    ]
    for field, value in fields:
        row = table.add_row().cells
        row[0].text = field
        row[1].text = value
    document.add_paragraph()


def _render_section_blocks(document: Document, section) -> None:
    if not section.blocks:
        for paragraph in [part.strip() for part in section.content.split("\n") if part.strip()]:
            body = document.add_paragraph(paragraph)
            body.paragraph_format.first_line_indent = Inches(0.25)
        return

    for block in section.blocks:
        if block.block_type == "image" and block.asset_path and Path(block.asset_path).exists():
            try:
                document.add_picture(block.asset_path, width=Inches(5.8))
                figure = document.paragraphs[-1]
                figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if block.caption:
                    caption = document.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = caption.add_run(block.caption)
                    run.italic = True
                    run.font.size = Pt(10)
            except Exception:
                fallback = document.add_paragraph(f"[Image unavailable: {block.alt_text or block.text or block.asset_path}]")
                fallback.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if block.block_type == "table" and block.rows:
            row_count = len(block.rows)
            col_count = max(len(row) for row in block.rows)
            table = document.add_table(rows=row_count, cols=col_count)
            table.style = "Table Grid"
            for row_index, row in enumerate(block.rows):
                for col_index, value in enumerate(row):
                    table.rows[row_index].cells[col_index].text = value
            document.add_paragraph()
            continue

        text = (block.text or "").strip()
        if not text:
            continue

        if block.block_type == "caption":
            caption = document.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption.add_run(text)
            run.italic = True
            run.font.size = Pt(10)
            continue

        if block.block_type == "callout":
            callout = document.add_paragraph()
            callout.paragraph_format.left_indent = Inches(0.35)
            callout.paragraph_format.right_indent = Inches(0.2)
            run = callout.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x3F, 0x5D, 0x73)
            continue

        body = document.add_paragraph(text)
        body.paragraph_format.first_line_indent = Inches(0.25)


def export_docx(book: Book, updates: List[WrittenUpdate], output_docx: str) -> Tuple[bool, str, Path]:
    output_path = Path(output_docx)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _style_document(document)
    _add_title_page(document, book.book_title)
    _add_export_summary(document, book, updates)

    updates_by_chapter = {}
    for update in assign_subsection_ids(updates):
        updates_by_chapter.setdefault(update.chapter_id, []).append(update)

    for chapter in book.chapters:
        title = chapter.title
        if not title.lower().startswith("chapter"):
            title = f"Chapter {chapter.chapter_id}: {title}"
        document.add_heading(title, level=1)
        intro = document.add_paragraph()
        intro_run = intro.add_run(
            f"This chapter contains {len(chapter.sections)} sections and "
            f"{len(updates_by_chapter.get(chapter.chapter_id, []))} accepted updates in this export."
        )
        intro_run.italic = True

        for section in chapter.sections:
            document.add_heading(f"{section.section_id} {section.title}", level=2)
            _render_section_blocks(document, section)

        chapter_updates = sorted(
            updates_by_chapter.get(chapter.chapter_id, []),
            key=lambda item: (item.section_id, item.proposed_subsection_id or ""),
        )
        if not chapter_updates:
            continue

        document.add_heading("Recent Advances", level=2)
        for update in chapter_updates:
            title_line = update.title or (update.text.split("\n\n")[0].strip() if update.text else "Generated Update")
            document.add_heading(f"{update.proposed_subsection_id} {title_line}", level=3)
            mapping_note = document.add_paragraph()
            mapping_run = mapping_note.add_run(f"Mapped to Section {update.section_id}")
            mapping_run.italic = True
            mapping_run.font.color.rgb = RGBColor(0x3F, 0x5D, 0x73)
            for paragraph in [part.strip() for part in update.text.split("\n\n")[1:] if part.strip()]:
                body = document.add_paragraph(paragraph)
                body.paragraph_format.first_line_indent = Inches(0.25)
            for source_line in _source_lines(update):
                source_para = document.add_paragraph(style="List Bullet")
                source_run = source_para.add_run(source_line)
                source_run.font.size = Pt(10)
        document.add_paragraph()

    document.save(output_path)
    manifest = build_export_manifest(
        format_name="docx",
        source_markdown=None,
        output_file=str(output_path),
        engine="python-docx",
        success=True,
    )
    return True, "python-docx", write_export_manifest(str(output_path), manifest)
