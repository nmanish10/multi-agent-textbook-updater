from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from PIL import Image

from core.models import Book, Chapter, Section, WrittenUpdate
from ingestion.pipeline import load_book
from rendering.docx_exporter import export_docx


class DocxSupportTests(unittest.TestCase):
    def test_docx_export_writes_file_and_manifest(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            output_docx = str(Path(temp_dir) / "updated.docx")
            book = Book(
                book_title="DOCX Book",
                chapters=[
                    Chapter(
                        chapter_id="1",
                        title="Signal Processing",
                        content="Signals overview",
                        sections=[Section(section_id="1.1", title="Sampling", content="Sampling text.")],
                    )
                ],
            )
            updates = [
                WrittenUpdate(
                    chapter_id="1",
                    section_id="1.1",
                    title="Adaptive Sampling",
                    text="Adaptive Sampling\n\nSection 1.1 gains adaptive sampling methods.\n\nThis matters because students can compare static and dynamic sampling schemes.",
                    source="https://example.com",
                )
            ]

            ok, engine, manifest_path = export_docx(book, updates, output_docx)
            self.assertTrue(ok)
            self.assertEqual(engine, "python-docx")
            self.assertTrue(Path(output_docx).exists())
            self.assertTrue(manifest_path.exists())
            exported = Document(output_docx)
            all_text = "\n".join(p.text for p in exported.paragraphs if p.text.strip())
            self.assertIn("Updated Textbook Export", all_text)
            self.assertIn("Export Summary", all_text)
            self.assertIn("Recent Advances", all_text)

    def test_docx_ingestion_parses_headings(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "source.docx"
            document = Document()
            document.add_heading("Networks Textbook", level=0)
            document.add_heading("Chapter 1: Foundations", level=1)
            document.add_heading("1.1 Basics", level=2)
            document.add_paragraph("Foundational networking concepts and terminology.")
            document.save(source_path)

            book = load_book(str(source_path))
            self.assertGreaterEqual(len(book.chapters), 1)
            self.assertGreaterEqual(sum(len(ch.sections) for ch in book.chapters), 1)
            self.assertEqual(book.metadata.source_format, "docx")

    def test_docx_ingestion_and_export_preserve_images(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            image_path = temp_path / "figure.png"
            Image.new("RGB", (48, 30), color=(60, 120, 180)).save(image_path)

            source_path = temp_path / "illustrated.docx"
            document = Document()
            document.add_heading("Illustrated Book", level=0)
            document.add_heading("Chapter 1: Systems", level=1)
            document.add_paragraph("Introductory material before numbered sections.")
            document.add_picture(str(image_path))
            document.add_heading("1.1 Components", level=2)
            document.add_paragraph("Components interact through defined interfaces.")
            document.save(source_path)

            book = load_book(str(source_path))
            section_blocks = book.chapters[0].sections[0].blocks
            image_blocks = [block for block in section_blocks if block.block_type == "image"]
            self.assertGreaterEqual(len(image_blocks), 1)
            self.assertTrue(Path(image_blocks[0].asset_path).exists())
            self.assertGreaterEqual(book.parse_report.assets_detected, 1)

            output_docx = str(temp_path / "exported.docx")
            ok, _, _ = export_docx(book, [], output_docx)
            self.assertTrue(ok)
            self.assertTrue(Path(output_docx).exists())

    def test_docx_ingestion_preserves_tables(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "tables.docx"
            document = Document()
            document.add_heading("Structured Book", level=0)
            document.add_heading("Chapter 1: Systems", level=1)
            document.add_heading("1.1 Overview", level=2)
            table = document.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Component"
            table.rows[0].cells[1].text = "Purpose"
            table.rows[1].cells[0].text = "Parser"
            table.rows[1].cells[1].text = "Extracts headings"
            document.save(source_path)

            book = load_book(str(source_path))
            blocks = book.chapters[0].sections[0].blocks
            table_blocks = [block for block in blocks if block.block_type == "table"]
            self.assertEqual(len(table_blocks), 1)
            self.assertEqual(table_blocks[0].rows[0], ["Component", "Purpose"])


if __name__ == "__main__":
    unittest.main()
